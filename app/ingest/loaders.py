from pathlib import Path
from typing import List
from pypdf import PdfReader
import pandas as pd
import sqlite3
from docx import Document
import pdfplumber


def load_txt(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="ignore")


def load_pdf(file_path: Path) -> str:
    parts: List[str] = []
    with pdfplumber.open(str(file_path)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                parts.append(f"[PAGE {i}]\n{text}")
            else:
                parts.append(f"[PAGE {i}]\n")
    return "\n\n".join(parts).strip()



def load_docx(file_path: Path) -> str:
    doc = Document(str(file_path))
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def load_csv(file_path: Path, max_rows: int = 200) -> str:
    df = pd.read_csv(file_path)
    if max_rows and len(df) > max_rows:
        df = df.head(max_rows)
    
    return df.to_csv(index=False).strip()


def load_sqlite(file_path: Path, max_rows_per_table: int = 200) -> str:
    
    con = sqlite3.connect(str(file_path))
    try:
        cur = con.cursor()
        cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%';")
        tables = [r[0] for r in cur.fetchall()]

        blocks: List[str] = []
        blocks.append("SQLite database tables: " + (", ".join(tables) if tables else "(none)"))

        for t in tables:
            blocks.append(f"\n[TABLE {t}]")
            try:
                df = pd.read_sql_query(f"SELECT * FROM '{t}' LIMIT {int(max_rows_per_table)};", con)
                blocks.append(df.to_csv(index=False).strip() if not df.empty else "(empty)")
            except Exception as e:
                blocks.append(f"(failed to read table: {e})")

        return "\n".join(blocks).strip()
    finally:
        con.close()
